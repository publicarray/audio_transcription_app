#!/usr/bin/env python3
"""
Audio Transcription App with Speaker Diarization

This script processes audio files (WAV/MP3), transcribes multi-person conversations,
performs basic speaker diarization to identify different speakers,
and outputs the transcription to a Word document with color-coding for each speaker.

Libraries used:
- pydub: For audio processing and manipulation
- SpeechRecognition: For audio transcription
- python-docx: For creating Word documents
"""

import os
import math
import argparse
import tempfile
from typing import List, Tuple, Dict, Optional
from datetime import date

import speech_recognition as sr
from pydub import AudioSegment
from docx import Document
from docx.shared import RGBColor

# Define flags to track available dependencies
HAVE_WHISPER = False
HAVE_PYANNOTE = False
HAVE_TORCH = False

# Conditionally import torch
try:
    import torch
    HAVE_TORCH = True
except ImportError:
    pass

# Conditionally import faster-whisper
try:
    from faster_whisper import WhisperModel
    HAVE_WHISPER = True
except ImportError as e:
    print(f"Warning: faster-whisper not available: {e}")
    print("For enhanced transcription, install: pip install faster-whisper")

# Conditionally import pyannote.audio
if HAVE_TORCH:
    try:
        from pyannote.audio import Pipeline
        HAVE_PYANNOTE = True
    except ImportError as e:
        print(f"Warning: pyannote.audio not available: {e}")
        print("For speaker diarization, install: pip install pyannote.audio")
else:
    print("Warning: PyTorch not available. Speaker diarization with pyannote.audio requires PyTorch.")
    print("Install PyTorch first: pip install torch")
# Global variables to store model instances
model_size = "medium"
whisper_model = None
diarization_pipeline = None

def get_available_features() -> Dict[str, bool]:
    """
    Return a dictionary of available features based on installed dependencies.
    
    Returns:
        Dictionary with feature flags
    """
    return {
        "enhanced_transcription": HAVE_WHISPER,
        "speaker_diarization": HAVE_PYANNOTE and HAVE_TORCH,
        "basic_transcription": True,  # Always available through speech_recognition
    }
def get_whisper_model():
    """
    Lazy-load the WhisperModel when needed.
    
    Returns:
        WhisperModel or None if loading fails
    """
    global whisper_model
    if not HAVE_WHISPER:
        print("faster-whisper is not available. Using basic transcription instead.")
        return None
    
    print(f"Debug: HAVE_WHISPER is {HAVE_WHISPER}, attempting to initialize Whisper model")
        
    if whisper_model is None:
        try:
            # Import must be successful since HAVE_WHISPER is True
            from faster_whisper import WhisperModel
            print(f"Debug: WhisperModel successfully imported, initializing with model_size={model_size}")
            
            try:
                whisper_model = WhisperModel(model_size, device="cpu", compute_type="int8")
                print(f"Successfully loaded whisper model: {model_size}")
            except ImportError as ie:
                print(f"ImportError initializing WhisperModel: {ie}")
                print("This might be caused by incompatible ctranslate2 version. Try: pip install ctranslate2==4.5.0 faster-whisper --force-reinstall")
                whisper_model = None
            except RuntimeError as re:
                print(f"RuntimeError initializing WhisperModel: {re}")
                print("This might be caused by missing model files or CUDA issues. Using CPU version may help.")
                whisper_model = None
            except Exception as ex:
                print(f"Unexpected error initializing WhisperModel: {ex}")
                print(f"Error type: {type(ex).__name__}")
                whisper_model = None
        except Exception as e:
            print(f"Warning: Could not import or load whisper model: {e}")
            print(f"Detailed error: {type(e).__name__}: {str(e)}")
            whisper_model = None
    return whisper_model

def get_diarization_pipeline():
    """
    Lazy-load the PyAnnote pipeline when needed.
    
    Returns:
        Pipeline or None if loading fails
    """
    global diarization_pipeline
    
    if not HAVE_PYANNOTE or not HAVE_TORCH:
        missing = []
        if not HAVE_TORCH:
            missing.append("PyTorch (pip install torch)")
        if not HAVE_PYANNOTE:
            missing.append("pyannote.audio (pip install pyannote.audio)")
        
        print(f"Speaker diarization unavailable. Missing: {', '.join(missing)}")
        print("Using basic speaker identification instead.")
        return None
        
    if diarization_pipeline is None:
        try:
            # Check for HF token
            hf_token = os.environ.get("HF_TOKEN")
            if not hf_token:
                print("Warning: HF_TOKEN environment variable not set.")
                print("You need a HuggingFace token to use pyannote.audio.")
                print("1. Get a token at https://huggingface.co/settings/tokens")
                print("2. Accept the license at https://huggingface.co/pyannote/speaker-diarization")
                print("3. Set the HF_TOKEN environment variable")
                print("Using basic speaker identification instead.")
                return None
                
            # Import must be successful since HAVE_PYANNOTE is True
            from pyannote.audio import Pipeline
            diarization_pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization@2.1",
                use_auth_token=hf_token
            )
            # Use CUDA if available
            if torch.cuda.is_available():
                diarization_pipeline = diarization_pipeline.to(torch.device("cuda"))
            print("Loaded pyannote diarization model")
        except Exception as e:
            print(f"Warning: Could not load pyannote diarization model: {e}")
            print("Will fall back to basic speaker identification instead.")
            diarization_pipeline = None
    return diarization_pipeline

class AudioProcessor:
    """Handles audio file processing and segmentation."""
    
    def __init__(self, file_path: str):
        """
        Initialize the audio processor.
        
        Args:
            file_path: Path to the audio file.
        """
        self.file_path = file_path
        self.audio = self._load_audio()
        
    def _load_audio(self) -> AudioSegment:
        """
        Load audio file using pydub.
        
        Returns:
            AudioSegment object containing the audio data.
        """
        print(f"Loading audio file: {self.file_path}")
        file_ext = os.path.splitext(self.file_path)[1].lower()
        
        if file_ext == '.mp3':
            return AudioSegment.from_mp3(self.file_path)
        elif file_ext == '.wav':
            return AudioSegment.from_wav(self.file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Use .wav or .mp3")
    
    def split_audio(self, segment_length_ms: int = 10000, 
                    silence_threshold_db: int = -40, 
                    min_silence_len_ms: int = 500) -> List[AudioSegment]:
        """
        Split audio into segments, trying to split at silence for better speaker separation.
        
        Args:
            segment_length_ms: Maximum length of each segment in milliseconds
            silence_threshold_db: The silence threshold in dB
            min_silence_len_ms: Minimum length of silence to be considered a potential split point
            
        Returns:
            List of audio segments.
        """
        print("Splitting audio into segments...")
        
        segments = []
        total_length_ms = len(self.audio)
        
        # If audio is shorter than segment_length_ms, return it as a single segment
        if total_length_ms <= segment_length_ms:
            return [self.audio]
        
        # Calculate how many segments we need
        num_segments = math.ceil(total_length_ms / segment_length_ms)
        
        # Find silent points that could be good splitting points
        silent_points = []
        for i in range(1, num_segments):
            target_point = i * segment_length_ms
            search_start = max(0, target_point - 2000)
            search_end = min(total_length_ms, target_point + 2000)
            
            # Look for silence around the target point
            for j in range(search_start, search_end, 100):
                segment = self.audio[j:j+min_silence_len_ms]
                if segment.dBFS < silence_threshold_db:
                    silent_points.append(j + min_silence_len_ms // 2)
                    break
            else:
                # If no silence found, just use the target point
                silent_points.append(target_point)
        
        # Split the audio at the identified points
        start_point = 0
        for point in silent_points:
            segments.append(self.audio[start_point:point])
            start_point = point
        
        # Add the last segment
        segments.append(self.audio[start_point:])
        
        print(f"Audio split into {len(segments)} segments")
        return segments
    
    def segment_to_wav(self, segment: AudioSegment, temp_path: str) -> str:
        """
        Convert an audio segment to a temporary WAV file for processing with SpeechRecognition.
        
        Args:
            segment: Audio segment to convert
            temp_path: Path to save the temporary WAV file
            
        Returns:
            Path to the temporary WAV file.
        """
        segment.export(temp_path, format="wav")
        return temp_path


class Transcriber:
    """Handles audio transcription and speaker diarization."""
    
    def __init__(self, num_speakers: int = 3):
        """
        Initialize the transcriber.
        
        Args:
            num_speakers: Expected number of speakers (used for fallback method).
        """
        self.recognizer = sr.Recognizer()
        self.num_speakers = num_speakers
        # We'll check for availability when the methods are actually called

        self.available_features = get_available_features()
        self.use_diarization = self.available_features["speaker_diarization"]  # Will attempt to use diarization if available
    
    def transcribe_segment(self, segment_path: str) -> str:
        """
        Transcribe an audio segment using faster-whisper.
        
        Args:
            segment_path: Path to the audio segment file.
            
        Returns:
            Transcribed text.
        """
        # First try faster-whisper if available
        if self.available_features["enhanced_transcription"]:
            try:
                # Get the model and attempt to use faster-whisper for transcription
                whisper_model = get_whisper_model()
                if whisper_model:
                    segments, _ = whisper_model.transcribe(segment_path, language="en")
                    text = " ".join([segment.text for segment in segments])
                    return text.strip()
            except Exception as e:
                print(f"Error with faster-whisper: {e}")
                print("Falling back to Google Speech Recognition")
        else:
            print("Enhanced transcription not available, using Google Speech Recognition")
            
        # Fall back to Google Speech Recognition
        try:
            with sr.AudioFile(segment_path) as source:
                audio_data = self.recognizer.record(source)
                text = self.recognizer.recognize_google(audio_data)
                return text
        except sr.UnknownValueError:
            return "[Inaudible]"
        except sr.RequestError:
            return "[Error: Could not request results from speech recognition service]"
    def identify_speakers(self, segments: List[AudioSegment]) -> List[Tuple[str, int, float, float]]:
        """
        Identify speakers using pyannote.audio for diarization.
        Falls back to basic heuristics if diarization fails.
        
        Args:
            segments: List of audio segments.
            
        Returns:
            List of (transcribed_text, speaker_id, start_time, end_time) tuples.
        """
        # First, try to use pyannote.audio for speaker diarization
        if self.use_diarization:
            pipeline = get_diarization_pipeline()
            if pipeline:
                try:
                    print("Performing speaker diarization with pyannote.audio...")
                    return self._identify_speakers_with_pyannote(segments)
                except Exception as e:
                    print(f"Diarization failed: {e}")
                    print("Falling back to basic speaker identification...")
            else:
                print("Diarization pipeline not available, using basic speaker identification")
                
        # Fall back to basic speaker identification if diarization fails or is unavailable
        return self._identify_speakers_basic(segments)
    
    def _identify_speakers_with_pyannote(self, segments: List[AudioSegment]) -> List[Tuple[str, int, float, float]]:
        """
        Identify speakers using pyannote.audio for diarization and faster-whisper for transcription.
        
        Args:
            segments: List of audio segments.
            
        Returns:
            List of (transcribed_text, speaker_id, start_time, end_time) tuples.
        """
        results = []
        
        # Combine all segments into a single audio file for diarization
        combined_audio = AudioSegment.empty()
        segment_offsets = []
        
        for segment in segments:
            segment_offsets.append(len(combined_audio) / 1000.0)  # Convert ms to seconds
            combined_audio += segment
        
        # Create a temporary WAV file for the combined audio
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
            combined_path = temp_file.name
            combined_audio.export(combined_path, format="wav")
            
            # Get the models
            pipeline = get_diarization_pipeline()
            whisper_model = get_whisper_model()
            
            if not pipeline:
                raise Exception("Diarization pipeline not available")
                
            if not whisper_model:
                raise Exception("Whisper model not available")
            
            # Perform diarization using pyannote.audio
            diarization = pipeline(combined_path)
            
            # Transcribe using faster-whisper
            whisper_segments, _ = whisper_model.transcribe(combined_path, word_timestamps=True)
            whisper_segments = list(whisper_segments)
            
            # Match whisper segments with speaker information
            speaker_map = {}  # Map speaker label strings to integer IDs
            
            for segment in whisper_segments:
                # Find the most common speaker for this segment's time range
                segment_start = segment.start
                segment_end = segment.end
                
                speaker_votes = {}
                # Find which speaker is speaking during this segment
                for turn, _, speaker in diarization.itertracks(yield_label=True):
                    # Check if the turn overlaps with the segment
                    if not (segment_end <= turn.start or segment_start >= turn.end):
                        # There is overlap
                        speaker_votes[speaker] = speaker_votes.get(speaker, 0) + 1
                
                # Get most frequent speaker for this segment
                if speaker_votes:
                    most_common_speaker = max(speaker_votes.items(), key=lambda x: x[1])[0]
                    
                    # Convert speaker string to integer ID
                    if most_common_speaker not in speaker_map:
                        speaker_map[most_common_speaker] = len(speaker_map)
                    
                    speaker_id = speaker_map[most_common_speaker]
                else:
                    speaker_id = 0
                
                # Add the segment to results
                if segment.text and segment.text.strip() != "[Inaudible]":
                    results.append((segment.text.strip(), speaker_id, segment.start, segment.end))
            
            # Clean up the temporary file
            os.unlink(combined_path)
            
        return results
    
    def _identify_speakers_basic(self, segments: List[AudioSegment]) -> List[Tuple[str, int, float, float]]:
        """
        Identify speakers using a basic heuristic approach.
        This simulates speaker diarization by assigning speakers based on segment characteristics.
        
        Args:
            segments: List of audio segments.
            
        Returns:
            List of (transcribed_text, speaker_id, start_time, end_time) tuples.
        """
        print("Performing transcription with basic speaker identification...")
        results = []
        
        # Analyze audio features to group by potential speakers
        segment_features = []
        for i, segment in enumerate(segments):
            # Use volume (dBFS) and other properties as simple features
            segment_features.append({
                'index': i,
                'dBFS': segment.dBFS,
                'duration': len(segment),
                'segment': segment
            })
        
        # Sort segments by volume as a simple way to cluster potential speakers
        segment_features.sort(key=lambda x: x['dBFS'])
        
        # Assign speaker IDs (0, 1, 2) based on sorted volume
        speaker_assignments = {}
        num_segments = len(segment_features)
        
        for i, feature in enumerate(segment_features):
            segment_idx = feature['index']
            # Divide segments into speaker groups based on volume
            speaker_id = min(int(i * self.num_speakers / num_segments), self.num_speakers - 1)
            speaker_assignments[segment_idx] = speaker_id
        
        # Transcribe each segment and assign speaker
        current_time = 0.0  # Track start time of each segment
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
            temp_path = temp_file.name
            
            for i, segment in enumerate(segments):
                # Export segment directly to WAV file
                segment.export(temp_path, format="wav")
                
                text = self.transcribe_segment(temp_path)
                speaker_id = speaker_assignments.get(i, 0)
                
                # Calculate start and end times
                segment_duration_sec = len(segment) / 1000.0  # Convert ms to seconds
                segment_start = current_time
                segment_end = current_time + segment_duration_sec
                current_time = segment_end
                
                if text and text != "[Inaudible]":
                    results.append((text, speaker_id, segment_start, segment_end))
            
            # Clean up
            os.unlink(temp_path)
            
        return results


class DocumentCreator:
    """Creates a Word document with color-coded transcription."""
    
    def __init__(self):
        """Initialize the document creator."""
        self.document = Document()
        self.speaker_colors = [
            RGBColor(0, 0, 200),  # Blue
            RGBColor(200, 0, 0),  # Red
            RGBColor(0, 150, 0),  # Green
        ]
    
    def add_transcription(self, transcription_data: List[Tuple[str, int, float, float]]):
        """
        Add the transcribed text to the document with color-coding for each speaker
        and timestamp information.
        
        Args:
            transcription_data: List of (transcribed_text, speaker_id, start_time, end_time) tuples.
        """
        print("Creating Word document with color-coded speakers...")
        
        # Add a title
        formatted_date = date.today().strftime("%d/%m/%Y")
        self.document.add_heading('Transcription ' + formatted_date, 0)
        
        # Create a legend for speaker colors
        legend = self.document.add_paragraph('Speakers: ')
        for i in range(len(self.speaker_colors)):
            speaker_run = legend.add_run(f"Speaker {i+1} ")
            speaker_run.font.color.rgb = self.speaker_colors[i]
        
        self.document.add_paragraph()  # Add a blank line
        
        # Add the transcription
        for text, speaker_id, start_time, end_time in transcription_data:
            speaker_num = speaker_id + 1  # Convert 0-based to 1-based for display
            p = self.document.add_paragraph()
            
            # Add timestamp
            time_str = f"[{self._format_time(start_time)}-{self._format_time(end_time)}] "
            time_label = p.add_run(time_str)
            time_label.italic = True
            
            # Add speaker label
            speaker_label = p.add_run(f"Speaker {speaker_num}: ")
            speaker_label.bold = True
            speaker_label.font.color.rgb = self.speaker_colors[speaker_id % len(self.speaker_colors)]
            
            # Add transcribed text
            text_run = p.add_run(text)
            text_run.font.color.rgb = self.speaker_colors[speaker_id % len(self.speaker_colors)]
    
    def _format_time(self, seconds: float) -> str:
        """
        Format seconds to MM:SS format.
        
        Args:
            seconds: Time in seconds.
            
        Returns:
            Formatted time string in MM:SS format.
        """
        minutes = int(seconds // 60)
        seconds = int(seconds % 60)
        return f"{minutes:02d}:{seconds:02d}"
    
    def save_document(self, output_path: str):
        """Save the document to the specified path.
        
        Args:
            output_path: Path to save the Word document.
        """
        self.document.save(output_path)
        print(f"Transcription saved to {output_path}")


def main():
    """Main function to run the audio transcription app."""
    # Parse command line arguments
    parser = argparse.ArgumentParser(description="Audio Transcription App with Speaker Diarization")
    parser.add_argument("input_file", help="Path to the input audio file (WAV or MP3)")
    parser.add_argument("--output", "-o", default="transcription.docx", 
                        help="Path to the output Word document (default: transcription.docx)")
    parser.add_argument("--speakers", "-s", type=int, default=3, 
                        help="Maximum number of speakers to identify (default: 3)")
    args = parser.parse_args()
    
    # Check available features and show info
    features = get_available_features()
    print("\nAvailable features:")
    print(f"  Enhanced transcription: {'Yes' if features['enhanced_transcription'] else 'No'}")
    print(f"  Speaker diarization: {'Yes' if features['speaker_diarization'] else 'No'}")
    print(f"  Basic transcription: {'Yes' if features['basic_transcription'] else 'No'}\n")
    
    if not features['enhanced_transcription']:
        print("Note: For better transcription quality, install faster-whisper:")
        print("  pip install faster-whisper")
        
    if not features['speaker_diarization']:
        print("Note: For accurate speaker identification, install pyannote.audio and torch:")
        print("  pip install torch pyannote.audio")
        print("  You'll also need a HuggingFace token set as the HF_TOKEN environment variable.")
        
    # Process the audio file
    try:
        # Step 1: Process the audio file
        processor = AudioProcessor(args.input_file)
        
        # Step 2: Split the audio into segments
        segments = processor.split_audio()
        
        # Step 3: Transcribe the audio with speaker identification
        transcriber = Transcriber(num_speakers=args.speakers)
        transcription_data = transcriber.identify_speakers(segments)
        
        # Step 4: Create the Word document
        doc_creator = DocumentCreator()
        doc_creator.add_transcription(transcription_data)
        doc_creator.save_document(args.output)
        
        print("Transcription complete!")
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()

